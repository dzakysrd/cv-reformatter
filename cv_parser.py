"""
cv_parser.py
Ekstraksi teks + heuristik segmentasi CV (tanpa AI/API berbayar).

Alur:
1. extract_blocks()  -> ubah file .docx/.pdf jadi list "blocks" (paragraf + info gaya tulisan)
2. segment_sections() -> kelompokkan blocks jadi bagian2 (kontak, pengalaman, pendidikan, dst)
   berdasarkan kata kunci bilingual (keywords.json)
3. parse_contact(), parse_entries(), parse_flat_list() -> ubah tiap bagian jadi data terstruktur
   (dict/list) yang siap dipakai cv_builder.py

Catatan penting: karena tidak pakai AI, hasil ekstraksi bersifat "best effort" berbasis
pola umum CV (heading pendek, tanggal, bullet). CV dengan format sangat tidak biasa mungkin
perlu dikoreksi manual di layar review pada app.py.
"""

import json
import os
import re
from difflib import SequenceMatcher

from docx import Document

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None


BULLET_CHARS = [
    "\u2022", "\u25E6", "\u25AA", "\u25CF", "\u00B7", "-", "*", "\u2023", "\u25CB",
    "\uf0b7", "\uf0a7", "\uf0d8", "\uf0fc", "\uf020",  # glyph bullet umum dari font Symbol/Wingdings
    # (sering muncul saat teks PDF hasil export Word diekstrak)
]

EMAIL_RE = re.compile(r"[\w\.\-+]+@[\w\-]+\.[\w\.\-]+")
PHONE_RE = re.compile(r"(\+?\d[\d\-\s\(\)]{7,}\d)")
LINKEDIN_RE = re.compile(r"(linkedin\.com/\S+)", re.IGNORECASE)
URL_RE = re.compile(r"((https?://)?(www\.)?[\w\-]+\.[a-z]{2,}(/\S*)?)", re.IGNORECASE)

MONTHS = (
    "jan(uari)?|feb(ruari)?|mar(et)?|apr(il)?|mei|jun[i]?|jul[i]?|agu(stus)?|"
    "sep(tember)?|okt(ober)?|nov(ember)?|des(ember)?|"
    "january|february|march|april|may|june|july|august|september|october|november|december"
)
DATE_RANGE_RE = re.compile(
    rf"((({MONTHS})\.?\s*)?\b(19|20)\d{{2}}\b)\s*[-–—~]{{1,2}}\s*"
    rf"((({MONTHS})\.?\s*)?\b(19|20)\d{{2}}\b|sekarang|present|now|current|saat ini)",
    re.IGNORECASE,
)
YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")


def _load_keywords(path=None):
    path = path or os.path.join(os.path.dirname(__file__), "keywords.json")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _normalize(text):
    return re.sub(r"[\s:.\-–—]+$", "", text.strip().lower()).strip()


def _is_bullet_line(text):
    stripped = text.strip()
    return any(stripped.startswith(ch + " ") or stripped == ch for ch in BULLET_CHARS)


def _strip_bullet(text):
    stripped = text.strip()
    for ch in BULLET_CHARS:
        if stripped.startswith(ch + " "):
            return stripped[len(ch) :].strip()
        if stripped.startswith(ch):
            return stripped[len(ch) :].strip()
    return stripped


# ---------------------------------------------------------------------------
# Ekstraksi mentah dari file -> list of blocks
# block = {"text": str, "bold": bool, "size": float|None, "is_bullet": bool}
# ---------------------------------------------------------------------------

def _blocks_from_docx(path):
    doc = Document(path)
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        bold = False
        size = None
        if para.runs:
            bold = any(r.bold for r in para.runs if r.bold is not None) or bool(para.runs[0].bold)
            sizes = [r.font.size.pt for r in para.runs if r.font and r.font.size]
            if sizes:
                size = max(sizes)
        is_list_style = bool(para.style and para.style.name and "list" in para.style.name.lower())
        blocks.append(
            {
                "text": text,
                "bold": bold,
                "size": size,
                "is_bullet": is_list_style or _is_bullet_line(text),
            }
        )
    # juga baca isi tabel (banyak CV pakai tabel untuk layout 2 kolom)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    bold = bool(para.runs and para.runs[0].bold)
                    blocks.append({"text": text, "bold": bold, "size": None, "is_bullet": _is_bullet_line(text)})
    return blocks


def _blocks_from_pdf(path):
    if pdfplumber is None:
        raise RuntimeError("pdfplumber belum terinstall. Jalankan: pip install pdfplumber")
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                blocks.append({"text": line, "bold": False, "size": None, "is_bullet": _is_bullet_line(line)})
    return blocks


def extract_blocks(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _blocks_from_docx(path)
    if ext == ".pdf":
        return _blocks_from_pdf(path)
    raise ValueError(f"Format file belum didukung: {ext} (pakai .docx atau .pdf)")


# ---------------------------------------------------------------------------
# Deteksi heading & segmentasi
# ---------------------------------------------------------------------------

def match_section(text, keywords):
    """Kembalikan (canonical_id, lang) kalau `text` cocok sebagai judul section, else (None, None)."""
    norm = _normalize(text)
    if not norm or len(norm.split()) > 6:
        return None, None
    for canonical, by_lang in keywords.items():
        for lang, variants in by_lang.items():
            for kw in variants:
                if norm == kw:
                    return canonical, lang
                ratio = SequenceMatcher(None, norm, kw).ratio()
                if ratio > 0.82:
                    return canonical, lang
                if kw in norm and len(kw) / max(len(norm), 1) > 0.55:
                    return canonical, lang
    return None, None


def segment_sections(blocks, keywords):
    """
    Pisahkan blocks jadi:
      - header_blocks: blocks sebelum heading pertama (biasanya nama & kontak)
      - sections: {canonical_id: [blocks]}
      - order: urutan canonical_id sesuai kemunculan
      - heading_text: {canonical_id: teks asli heading di dokumen ini}
      - lang_votes: list bahasa yang match, buat nebak bahasa dominan dokumen
    """
    header_blocks = []
    sections = {}
    order = []
    heading_text = {}
    lang_votes = []

    current = None
    for b in blocks:
        canonical, lang = (None, None) if b["is_bullet"] else match_section(b["text"], keywords)
        if canonical:
            current = canonical
            if canonical not in sections:
                sections[canonical] = []
                order.append(canonical)
                heading_text[canonical] = b["text"].strip()
            lang_votes.append(lang)
            continue
        if current is None:
            header_blocks.append(b)
        else:
            sections[canonical or current].append(b)

    return {
        "header_blocks": header_blocks,
        "sections": sections,
        "order": order,
        "heading_text": heading_text,
        "lang_votes": lang_votes,
    }


# ---------------------------------------------------------------------------
# Parsing tiap bagian jadi data terstruktur
# ---------------------------------------------------------------------------

def parse_contact(header_blocks):
    name = ""
    email = ""
    phone = ""
    linkedin = ""
    extra_lines = []

    for b in header_blocks:
        text = b["text"].strip()
        if not text:
            continue

        remaining = text
        found_something = False

        m_email = EMAIL_RE.search(remaining)
        if m_email and not email:
            email = m_email.group(0)
            remaining = remaining.replace(m_email.group(0), " ")
            found_something = True

        m_linkedin = LINKEDIN_RE.search(remaining)
        if m_linkedin and not linkedin:
            linkedin = m_linkedin.group(0)
            remaining = remaining.replace(m_linkedin.group(0), " ")
            found_something = True

        m_phone = PHONE_RE.search(remaining)
        if m_phone and not phone and len(re.sub(r"\D", "", m_phone.group(0))) >= 8:
            phone = m_phone.group(0).strip()
            remaining = remaining.replace(m_phone.group(0), " ")
            found_something = True

        remainder = re.sub(r"\s*\|\s*(\|\s*)*", " | ", remaining).strip(" |,-–—")

        if found_something:
            if remainder:
                extra_lines.append(remainder)
        elif not name:
            name = text
        else:
            extra_lines.append(text)

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "location": " | ".join(extra_lines[:2]),
    }


def _split_header_line(line):
    """Coba pecah 1 baris 'Jabatan - Perusahaan | Jan 2020 - Sekarang' jadi (judul, sub, tanggal)."""
    date_match = DATE_RANGE_RE.search(line)
    date_text = date_match.group(0) if date_match else ""
    remainder = line
    if date_match:
        remainder = line[: date_match.start()] + line[date_match.end() :]
        remainder = re.sub(r"\(\s*\)", "", remainder)  # bersihkan "()" kosong sisa tanggal
        remainder = re.sub(r"\s{2,}", " ", remainder).strip(" |,-–—")
    parts = re.split(r"\s*[|–—-]\s*", remainder)
    parts = [p.strip() for p in parts if p.strip()]
    title = parts[0] if parts else remainder
    subtitle = " - ".join(parts[1:]) if len(parts) > 1 else ""
    return title, subtitle, date_text


def parse_entries(section_blocks):
    """
    Pecah isi 1 section (misal Pengalaman Kerja) jadi list of entries:
    [{"title":..., "subtitle":..., "date":..., "bullets":[...]}, ...]
    """
    entries = []
    current = None

    for b in section_blocks:
        text = b["text"].strip()
        if not text:
            continue
        if b["is_bullet"]:
            if current is None:
                current = {"title": "", "subtitle": "", "date": "", "bullets": []}
                entries.append(current)
            current["bullets"].append(_strip_bullet(text))
            continue

        starts_new = current is None or current["bullets"] or (current["title"] and current["subtitle"] and current["date"])
        if starts_new:
            title, subtitle, date = _split_header_line(text)
            current = {"title": title, "subtitle": subtitle, "date": date, "bullets": []}
            entries.append(current)
        else:
            # baris ke-2/3 dari header entry yang sama (misal perusahaan di baris terpisah dari tanggal)
            date_match = DATE_RANGE_RE.search(text)
            if date_match and not current["date"]:
                current["date"] = date_match.group(0)
                rest = (text[: date_match.start()] + text[date_match.end() :]).strip(" |,-–—")
                if rest and not current["subtitle"]:
                    current["subtitle"] = rest
            elif not current["subtitle"]:
                current["subtitle"] = text
            else:
                current["subtitle"] += " | " + text

    return entries


def parse_flat_list(section_blocks):
    """Buat section sederhana (skills/languages/awards) jadi list string bersih."""
    items = []
    for b in section_blocks:
        text = _strip_bullet(b["text"].strip())
        if not text:
            continue
        if b["is_bullet"] or len(section_blocks) > 3:
            items.append(text)
        else:
            # satu paragraf panjang dipisah koma / titik koma / slash
            parts = re.split(r"[,;/]|(?:\s{2,})", text)
            items.extend([p.strip() for p in parts if p.strip()])
    # buang duplikat sambil jaga urutan
    seen = set()
    result = []
    for it in items:
        key = it.lower()
        if key not in seen:
            seen.add(key)
            result.append(it)
    return result


LIST_SECTIONS = {"skills", "languages", "awards", "references"}
ENTRY_SECTIONS = {"experience", "education", "certifications", "organizations", "projects"}


def parse_cv(path, keywords=None):
    """Fungsi utama: path file CV -> dict data terstruktur siap dipakai builder."""
    keywords = keywords or _load_keywords()
    blocks = extract_blocks(path)
    seg = segment_sections(blocks, keywords)

    data = {"contact": parse_contact(seg["header_blocks"]), "order": seg["order"], "sections": {}}

    for canonical, sec_blocks in seg["sections"].items():
        if canonical in ENTRY_SECTIONS:
            data["sections"][canonical] = parse_entries(sec_blocks)
        elif canonical in LIST_SECTIONS:
            data["sections"][canonical] = parse_flat_list(sec_blocks)
        elif canonical == "summary":
            data["sections"][canonical] = " ".join(b["text"].strip() for b in sec_blocks if b["text"].strip())
        else:
            data["sections"][canonical] = parse_flat_list(sec_blocks)

    return data


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Pemakaian: python cv_parser.py path/ke/cv.docx")
        sys.exit(1)
    result = parse_cv(sys.argv[1])
    print(json.dumps(result, indent=2, ensure_ascii=False))
