"""
cv_builder.py
Bangun CV baru dari data hasil ekstraksi (cv_parser.parse_cv), mengikuti STRUKTUR & GAYA
visual dari file template (contoh CV yang formatnya ingin ditiru).

Karena template di sini adalah "contoh CV biasa" (bukan file dengan tag {{placeholder}}),
pendekatannya:
  1. analyze_template()  -> baca template.docx, kenali urutan section + contoh gaya tulisan
     tiap bagian (font, ukuran, bold/italic, bullet, margin halaman, dst).
  2. build_cv()           -> bikin dokumen docx BARU dari nol, lalu isi tiap section pakai
     data CV sumber, dengan gaya tulisan meniru sampel dari template.

Pendekatan ini lebih aman/stabil dibanding "menyuntik" isi ke file template asli (yang gampang
merusak struktur XML docx), meski detail layout yang sangat unik (kolom ganda, tabel dekoratif,
dsb.) belum tentu tertiru 100%.
"""

from docx import Document
from docx.shared import Pt, RGBColor

from cv_parser import (
    ENTRY_SECTIONS,
    LIST_SECTIONS,
    BULLET_CHARS,
    match_section,
    _load_keywords,
)

DEFAULT_LABELS = {
    "id": {
        "summary": "Ringkasan",
        "experience": "Pengalaman Kerja",
        "education": "Pendidikan",
        "skills": "Keahlian",
        "certifications": "Sertifikasi",
        "languages": "Bahasa",
        "organizations": "Organisasi",
        "projects": "Proyek",
        "awards": "Penghargaan",
        "references": "Referensi",
    },
    "en": {
        "summary": "Summary",
        "experience": "Experience",
        "education": "Education",
        "skills": "Skills",
        "certifications": "Certifications",
        "languages": "Languages",
        "organizations": "Organizations",
        "projects": "Projects",
        "awards": "Awards",
        "references": "References",
    },
}


def _style_from_para(para):
    run = para.runs[0] if para.runs else None
    style = {
        "font_name": None,
        "size": None,
        "bold": False,
        "italic": False,
        "underline": False,
        "color": None,
        "alignment": para.alignment,
    }
    if run is not None:
        style["font_name"] = run.font.name
        if run.font.size:
            style["size"] = run.font.size.pt
        style["bold"] = bool(run.bold)
        style["italic"] = bool(run.italic)
        style["underline"] = bool(run.underline)
        if run.font.color and run.font.color.type is not None and run.font.color.rgb:
            style["color"] = str(run.font.color.rgb)
    return style


def _is_bullet_para(para):
    style_name = (para.style.name or "").lower() if para.style else ""
    if "list" in style_name or "bullet" in style_name:
        return True, None
    try:
        if para._p.pPr is not None and para._p.pPr.numPr is not None:
            return True, None
    except Exception:
        pass
    text = para.text.strip()
    for ch in BULLET_CHARS:
        if text.startswith(ch + " ") or text == ch:
            return True, ch
    return False, None


def analyze_template(template_path, keywords=None):
    keywords = keywords or _load_keywords()
    doc = Document(template_path)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    normal_style = doc.styles["Normal"]
    default_font_name = normal_style.font.name or "Calibri"
    default_font_size = normal_style.font.size.pt if normal_style.font.size else 11

    sec0 = doc.sections[0]
    margins = {
        "top": sec0.top_margin,
        "bottom": sec0.bottom_margin,
        "left": sec0.left_margin,
        "right": sec0.right_margin,
    }

    header_paras = []
    section_paras = {}
    order = []
    heading_text = {}
    heading_style = None
    lang_votes = []
    current = None

    for p in paragraphs:
        text = p.text.strip()
        is_bul, _ = _is_bullet_para(p)
        canonical, lang = (None, None) if is_bul else match_section(text, keywords)
        if canonical:
            current = canonical
            if canonical not in section_paras:
                section_paras[canonical] = []
                order.append(canonical)
                heading_text[canonical] = text
                if heading_style is None:
                    heading_style = _style_from_para(p)
            lang_votes.append(lang)
            continue
        if current is None:
            header_paras.append(p)
        else:
            section_paras.setdefault(canonical or current, []).append(p)

    name_style = _style_from_para(header_paras[0]) if header_paras else {}
    contact_style = _style_from_para(header_paras[1]) if len(header_paras) > 1 else name_style

    body_style = {
        "font_name": default_font_name,
        "size": default_font_size,
        "bold": False,
        "italic": False,
        "underline": False,
        "color": None,
        "alignment": None,
    }
    bullet_style = dict(body_style)
    bullet_char = "•"
    entry_title_style = None
    entry_subtitle_style = None

    for canonical in ENTRY_SECTIONS:
        paras = section_paras.get(canonical)
        if not paras:
            continue
        non_bullet = [p for p in paras if not _is_bullet_para(p)[0]]
        bullets = [p for p in paras if _is_bullet_para(p)[0]]
        if non_bullet:
            entry_title_style = _style_from_para(non_bullet[0])
            if len(non_bullet) > 1:
                entry_subtitle_style = _style_from_para(non_bullet[1])
        if bullets:
            bullet_style = _style_from_para(bullets[0])
            _, bc = _is_bullet_para(bullets[0])
            if bc:
                bullet_char = bc
        break

    if entry_title_style is None:
        entry_title_style = heading_style or body_style
    if entry_subtitle_style is None:
        entry_subtitle_style = body_style

    skills_paras = section_paras.get("skills", [])
    skills_inline = (len(skills_paras) <= 2) if skills_paras else True

    lang_votes = [l for l in lang_votes if l]
    template_lang = max(set(lang_votes), key=lang_votes.count) if lang_votes else "id"

    return {
        "order": order,
        "heading_text": heading_text,
        "heading_style": heading_style or body_style,
        "name_style": name_style,
        "contact_style": contact_style,
        "entry_title_style": entry_title_style,
        "entry_subtitle_style": entry_subtitle_style,
        "bullet_style": bullet_style,
        "bullet_char": bullet_char,
        "body_style": body_style,
        "default_font": {"name": default_font_name, "size": default_font_size},
        "margins": margins,
        "skills_inline": skills_inline,
        "template_lang": template_lang if template_lang in DEFAULT_LABELS else "id",
    }


def _apply_run_style(run, style):
    if not style:
        return
    if style.get("font_name"):
        run.font.name = style["font_name"]
    if style.get("size"):
        run.font.size = Pt(style["size"])
    run.bold = bool(style.get("bold"))
    run.italic = bool(style.get("italic"))
    run.underline = bool(style.get("underline"))
    if style.get("color"):
        try:
            run.font.color.rgb = RGBColor.from_string(style["color"])
        except Exception:
            pass


def _add_styled_para(doc, text, style):
    p = doc.add_paragraph()
    if style and style.get("alignment") is not None:
        p.alignment = style["alignment"]
    run = p.add_run(text)
    _apply_run_style(run, style)
    return p


def _add_bullet_para(doc, text, style):
    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
        text = "• " + text
    run = p.add_run(text)
    _apply_run_style(run, style)
    return p


def build_cv(parsed_data, template_analysis, output_path):
    doc = Document()

    sec0 = doc.sections[0]
    m = template_analysis.get("margins", {})
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        key = attr.split("_")[0]
        if m.get(key) is not None:
            setattr(sec0, attr, m[key])

    normal = doc.styles["Normal"]
    if template_analysis["default_font"].get("name"):
        normal.font.name = template_analysis["default_font"]["name"]
    if template_analysis["default_font"].get("size"):
        normal.font.size = Pt(template_analysis["default_font"]["size"])

    contact = parsed_data.get("contact", {})
    if contact.get("name"):
        _add_styled_para(doc, contact["name"], template_analysis["name_style"])
    contact_bits = [
        c for c in [contact.get("phone"), contact.get("email"), contact.get("linkedin"), contact.get("location")] if c
    ]
    if contact_bits:
        _add_styled_para(doc, " | ".join(contact_bits), template_analysis["contact_style"])
    doc.add_paragraph()

    labels = DEFAULT_LABELS[template_analysis["template_lang"]]
    template_order = template_analysis["order"]
    source_sections = parsed_data.get("sections", {})
    final_order = list(template_order) + [s for s in parsed_data.get("order", []) if s not in template_order]

    for canonical in final_order:
        content = source_sections.get(canonical)
        if not content:
            continue

        heading = template_analysis["heading_text"].get(canonical) or labels.get(canonical, canonical.title())
        _add_styled_para(doc, heading, template_analysis["heading_style"])

        if canonical in ENTRY_SECTIONS:
            for entry in content:
                title_line = entry.get("title", "")
                if entry.get("subtitle"):
                    title_line = f"{title_line} — {entry['subtitle']}" if title_line else entry["subtitle"]
                if title_line:
                    _add_styled_para(doc, title_line, template_analysis["entry_title_style"])
                if entry.get("date"):
                    _add_styled_para(doc, entry["date"], template_analysis["entry_subtitle_style"])
                for bullet in entry.get("bullets", []):
                    _add_bullet_para(doc, bullet, template_analysis["bullet_style"])
        elif canonical == "summary":
            _add_styled_para(doc, content, template_analysis["body_style"])
        else:
            items = content if isinstance(content, list) else [content]
            if canonical == "skills" and template_analysis.get("skills_inline") and len(items) > 1:
                _add_styled_para(doc, ", ".join(items), template_analysis["body_style"])
            else:
                for it in items:
                    _add_bullet_para(doc, it, template_analysis["bullet_style"])

        doc.add_paragraph()

    doc.save(output_path)
    return output_path


def reformat_cv(template_path, source_cv_path, output_path, parsed_data=None, keywords=None):
    """Fungsi convenience end-to-end (dipakai app.py & test)."""
    from cv_parser import parse_cv

    keywords = keywords or _load_keywords()
    if parsed_data is None:
        parsed_data = parse_cv(source_cv_path, keywords)
    template_analysis = analyze_template(template_path, keywords)
    return build_cv(parsed_data, template_analysis, output_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 4:
        print("Pemakaian: python cv_builder.py template.docx cv_sumber.docx hasil_output.docx")
        sys.exit(1)
    out = reformat_cv(sys.argv[1], sys.argv[2], sys.argv[3])
    print(f"Selesai -> {out}")
